#!/usr/bin/env python3
"""Generate HW1a Word report for MNIST MLP classification."""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGURES = Path("C:/Users/muham/Desktop/Masters Spring Semster 20252026/CS 515/Deep-learning-CS515/HW1/outputs/figures")
OUTPUT = Path("C:/Users/muham/Desktop/Masters Spring Semster 20252026/CS 515/Deep-learning-CS515/HW1/HW1_Report.docx")

IMG_FULL_W = Inches(5.8)
IMG_FULL_H = Inches(3.6)
IMG_SIDE_W = Inches(2.8)
IMG_SIDE_H = Inches(1.75)
IMG_OVERLAY_W = Inches(5.2)
IMG_OVERLAY_H = Inches(3.2)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_figure(doc, img_path, caption, width=None, height=None):
    if width is None:
        width = IMG_FULL_W
    if not Path(img_path).exists():
        add_para(doc, f"[Figure not found: {img_path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if height:
        run.add_picture(str(img_path), width=width, height=height)
    else:
        run.add_picture(str(img_path), width=width)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.paragraph_format.space_after = Pt(10)


def add_side_by_side(doc, img1, img2, cap1, cap2):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (img, cap) in enumerate([(img1, cap1), (img2, cap2)]):
        cell = table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if Path(img).exists():
            run = p.add_run()
            run.add_picture(str(img), width=IMG_SIDE_W, height=IMG_SIDE_H)
        else:
            p.add_run(f"[Not found: {img}]")
        cap_p = cell.add_paragraph(cap)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.runs[0].italic = True
        cap_p.runs[0].font.size = Pt(8)
    doc.add_paragraph()


def add_results_table(doc, rows_data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows_data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return table


def main():
    doc = Document()

    # ----- Page margins -----
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # =========================================================
    # TITLE PAGE
    # =========================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("HW1a: MNIST Classification with MLP")
    tr.bold = True
    tr.font.size = Pt(22)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("CS 515 - Deep Learning")
    sr.font.size = Pt(14)
    sr.bold = True

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run("Spring Semester 2025-2026").font.size = Pt(12)

    doc.add_page_break()

    # =========================================================
    # 1. PROBLEM DEFINITION
    # =========================================================
    add_heading(doc, "1. Problem Definition", 1)
    add_para(doc, (
        "This homework addresses multi-class image classification on the MNIST handwritten digit dataset "
        "using a Multi-Layer Perceptron (MLP). The goal is to correctly assign each 28x28 grayscale image "
        "to one of 10 digit classes (0-9) by training a fully-connected neural network."
    ))
    add_para(doc, (
        "The study involves a systematic ablation analysis covering: MLP architecture depth and width, "
        "choice of activation function (ReLU vs. GELU), dropout regularization, batch normalization, "
        "L1/L2 weight regularization, learning rate selection, learning rate scheduling, and early stopping. "
        "All experiments use the Adam optimizer with cross-entropy loss unless otherwise specified."
    ))

    # =========================================================
    # 2. NEURAL NETWORK MODEL
    # =========================================================
    add_heading(doc, "2. Neural Network Model", 1)
    add_para(doc, (
        "The ConfigurableMLP is a fully-connected neural network implemented with torch.nn.Sequential "
        "and torch.nn.ModuleList. Each hidden block follows the order:"
    ))

    # Block diagram as bullet list
    for item in [
        "Linear (torch.nn.Linear) — affine transformation",
        "BatchNorm1d (torch.nn.BatchNorm1d) — optional, placed BEFORE activation",
        "Activation — ReLU or GELU (torch.nn.ReLU / torch.nn.GELU)",
        "Dropout (torch.nn.Dropout) — optional stochastic regularization",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)

    add_para(doc, (
        "The input tensor of shape (B, 1, 28, 28) is first flattened to (B, 784) using torch.nn.Flatten. "
        "After all hidden blocks, a final torch.nn.Linear maps the last hidden dimension to 10 logits. "
        "No softmax is applied in the forward pass; it is handled implicitly by torch.nn.CrossEntropyLoss."
    ))

    add_para(doc, "Baseline configuration:", bold=True)
    rows = [
        ("Input dimension", "784 (28x28 flattened)"),
        ("Hidden layers", "2 layers: [256, 128]"),
        ("Output classes", "10"),
        ("Activation", "ReLU"),
        ("Dropout", "p = 0.2"),
        ("Batch Normalization", "Enabled (before activation)"),
        ("Optimizer", "Adam, lr = 0.001"),
        ("Loss", "CrossEntropyLoss"),
    ]
    add_results_table(doc, rows, ["Parameter", "Value"])

    # =========================================================
    # 3. DATASET
    # =========================================================
    add_heading(doc, "3. Dataset", 1)
    add_para(doc, (
        "The MNIST dataset consists of 70,000 grayscale images of handwritten digits (0-9), "
        "each of size 28x28 pixels. It is a standard benchmark for image classification."
    ))

    rows = [
        ("Training set", "55,000 samples (split from 60,000 MNIST train)"),
        ("Validation set", "5,000 samples (split from 60,000 MNIST train)"),
        ("Test set", "10,000 samples (MNIST official test set)"),
        ("Input size", "28 x 28 = 784 features (flattened)"),
        ("Classes", "10 (digits 0-9)"),
        ("Normalization", "mean=0.1307, std=0.3081"),
        ("Batch size", "128"),
    ]
    add_results_table(doc, rows, ["Split", "Details"])

    add_para(doc, (
        "The 60,000 training images are split into 55,000 training and 5,000 validation samples using a "
        "fixed random seed (42) for reproducibility. Normalization uses the dataset-wide mean and standard "
        "deviation. No data augmentation is applied."
    ))

    # =========================================================
    # 4. TRAINING FRAMEWORK
    # =========================================================
    add_heading(doc, "4. Training Framework", 1)
    add_para(doc, (
        "All models are trained using the following framework:"
    ))

    for item in [
        "Optimizer: Adam (default) or SGD with momentum",
        "Loss: Cross-Entropy (nn.CrossEntropyLoss)",
        "L1 regularization: custom penalty added to the loss",
        "L2 regularization: via weight_decay in the Adam optimizer",
        "LR Schedulers: StepLR (step decay every 5 epochs, gamma=0.5) or ReduceLROnPlateau (factor=0.5, patience=2)",
        "Early stopping: patience=5 epochs on validation loss improvement (min_delta=0)",
        "Best model checkpoint: saved when validation loss improves",
        "Max epochs: 20 (unless early stopping triggers sooner)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)

    add_para(doc, (
        "At each epoch, training loss/accuracy and validation loss/accuracy are recorded. "
        "The best checkpoint (lowest validation loss) is saved and used for final evaluation. "
        "The training and validation curves are plotted for each experiment."
    ))

    doc.add_page_break()

    # =========================================================
    # 5. NUMERICAL RESULTS AND ABLATIONS
    # =========================================================
    add_heading(doc, "5. Numerical Results and Ablations", 1)

    # --- 5.1 Baseline ---
    add_heading(doc, "5.1 Baseline", 2)
    add_para(doc, (
        "The baseline model uses 2 hidden layers [256, 128] with ReLU activation, BatchNorm, "
        "dropout p=0.2, Adam optimizer (lr=1e-3), and early stopping. "
        "It achieves a best validation accuracy of 97.92% at epoch 6."
    ))
    add_side_by_side(
        doc,
        FIGURES / "baseline_loss.png", FIGURES / "baseline_acc.png",
        "Figure 1a: Baseline Training/Validation Loss", "Figure 1b: Baseline Training/Validation Accuracy"
    )

    # --- 5.2 Architecture ---
    add_heading(doc, "5.2 Impact of MLP Architecture (Depth and Width)", 2)
    add_para(doc, (
        "Four architectures were compared by varying the number and width of hidden layers, "
        "with all other hyperparameters fixed to baseline values."
    ))

    rows = [
        ("arch_1x128", "1", "[128]", "97.68%"),
        ("arch_2x256_128 (baseline)", "2", "[256, 128]", "97.92%"),
        ("arch_3x512_256_128", "3", "[512, 256, 128]", "98.28%"),
        ("arch_wide_512_256", "2", "[512, 256]", "98.00%"),
    ]
    add_results_table(doc, rows, ["Run", "Layers", "Hidden Dims", "Best Val Acc"])

    add_para(doc, (
        "Deeper and wider networks consistently outperform shallower ones. The 3-layer architecture "
        "[512, 256, 128] achieves the best validation accuracy (98.28%), demonstrating that additional "
        "capacity helps capture more complex features. A single hidden layer of 128 units underperforms, "
        "while a wider 2-layer network [512, 256] offers a good trade-off. However, excessively deep "
        "networks without skip connections may eventually plateau or overfit for MNIST."
    ))

    add_figure(doc, FIGURES / "architecture_val_loss_overlay.png",
               "Figure 2a: Validation Loss Comparison - Architecture Ablation", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)
    add_figure(doc, FIGURES / "architecture_val_acc_overlay.png",
               "Figure 2b: Validation Accuracy Comparison - Architecture Ablation", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)

    add_side_by_side(
        doc,
        FIGURES / "arch_1x128_loss.png", FIGURES / "arch_1x128_acc.png",
        "Figure 3a: 1x128 Loss", "Figure 3b: 1x128 Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "arch_2x256_128_loss.png", FIGURES / "arch_2x256_128_acc.png",
        "Figure 4a: 2x[256,128] Loss", "Figure 4b: 2x[256,128] Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "arch_3x512_256_128_loss.png", FIGURES / "arch_3x512_256_128_acc.png",
        "Figure 5a: 3x[512,256,128] Loss", "Figure 5b: 3x[512,256,128] Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "arch_wide_512_256_loss.png", FIGURES / "arch_wide_512_256_acc.png",
        "Figure 6a: 2x[512,256] Wide Loss", "Figure 6b: 2x[512,256] Wide Accuracy"
    )

    doc.add_page_break()

    # --- 5.3 Activation Functions ---
    add_heading(doc, "5.3 Impact of Activation Functions", 2)
    add_para(doc, (
        "Two activation functions were compared: ReLU (Rectified Linear Unit) and GELU "
        "(Gaussian Error Linear Unit). Both use the baseline [256, 128] architecture."
    ))

    rows = [
        ("act_relu (ReLU)", "97.92%", "6"),
        ("act_gelu (GELU)", "97.94%", "6"),
    ]
    add_results_table(doc, rows, ["Run", "Best Val Acc", "Best Epoch"])

    add_para(doc, (
        "ReLU and GELU perform virtually identically on MNIST, with GELU achieving a marginally "
        "higher validation accuracy (97.94% vs 97.92%). Both converge at the same epoch (6). "
        "For simple datasets like MNIST, the choice of activation function has minimal impact. "
        "GELU may show more pronounced advantages on larger, more complex datasets where its "
        "smooth probabilistic gating provides better gradient flow."
    ))

    add_figure(doc, FIGURES / "activation_val_loss_overlay.png",
               "Figure 7a: Validation Loss - Activation Comparison", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)
    add_figure(doc, FIGURES / "activation_val_acc_overlay.png",
               "Figure 7b: Validation Accuracy - Activation Comparison", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)

    add_side_by_side(
        doc,
        FIGURES / "act_relu_loss.png", FIGURES / "act_relu_acc.png",
        "Figure 8a: ReLU Loss", "Figure 8b: ReLU Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "act_gelu_loss.png", FIGURES / "act_gelu_acc.png",
        "Figure 9a: GELU Loss", "Figure 9b: GELU Accuracy"
    )

    doc.add_page_break()

    # --- 5.4 Dropout ---
    add_heading(doc, "5.4 Impact of Dropout", 2)
    add_para(doc, (
        "Dropout was analyzed at three levels: 0.0 (disabled), 0.2 (baseline), and 0.5 (high). "
        "Dropout randomly zeroes neuron activations during training to reduce overfitting."
    ))

    rows = [
        ("dropout_0p0", "0.0", "97.74%", "4"),
        ("dropout_0p2 (baseline)", "0.2", "97.92%", "6"),
        ("dropout_0p5", "0.5", "98.08%", "20"),
    ]
    add_results_table(doc, rows, ["Run", "Dropout p", "Best Val Acc", "Best Epoch"])

    add_para(doc, (
        "Dropout improves validation accuracy: with no dropout (p=0.0) the model converges quickly (epoch 4) "
        "but with lower accuracy (97.74%), indicating mild overfitting. At p=0.2, the model achieves 97.92% "
        "with a good balance of regularization and convergence speed. At p=0.5, the model reaches the highest "
        "validation accuracy (98.08%) but requires all 20 epochs to converge, as strong dropout slows "
        "training significantly. The training curves show a larger train-val gap without dropout, confirming "
        "overfitting. Dropout forces the network to learn more robust, redundant representations."
    ))

    add_figure(doc, FIGURES / "dropout_val_loss_overlay.png",
               "Figure 10a: Validation Loss - Dropout Comparison", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)
    add_figure(doc, FIGURES / "dropout_val_acc_overlay.png",
               "Figure 10b: Validation Accuracy - Dropout Comparison", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)

    add_side_by_side(
        doc,
        FIGURES / "dropout_0p0_loss.png", FIGURES / "dropout_0p0_acc.png",
        "Figure 11a: Dropout 0.0 Loss", "Figure 11b: Dropout 0.0 Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "dropout_0p2_loss.png", FIGURES / "dropout_0p2_acc.png",
        "Figure 12a: Dropout 0.2 Loss", "Figure 12b: Dropout 0.2 Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "dropout_0p5_loss.png", FIGURES / "dropout_0p5_acc.png",
        "Figure 13a: Dropout 0.5 Loss", "Figure 13b: Dropout 0.5 Accuracy"
    )

    doc.add_page_break()

    # --- 5.5 Training Framework ---
    add_heading(doc, "5.5 Impact of Training Framework", 2)
    add_para(doc, "Learning rate, LR scheduler, and early stopping were evaluated.", bold=False)

    add_heading(doc, "5.5.1 Learning Rate", 3)
    rows = [
        ("lr_1e2", "0.01", "none", "97.94%", "6"),
        ("lr_1e3 (baseline)", "0.001", "none", "97.92%", "6"),
        ("lr_3e4", "0.0003", "none", "98.04%", "15"),
    ]
    add_results_table(doc, rows, ["Run", "LR", "Scheduler", "Best Val Acc", "Best Epoch"])

    add_para(doc, (
        "A smaller learning rate (3e-4) achieves the best accuracy (98.04%) among the fixed-LR runs, "
        "though it converges more slowly (epoch 15). The largest LR (1e-2) converges fast but with "
        "slightly lower accuracy (97.94%). The baseline lr=1e-3 strikes a balance between speed and "
        "performance. For Adam, a moderately small learning rate is generally preferred."
    ))

    add_figure(doc, FIGURES / "training_val_loss_overlay.png",
               "Figure 14a: Validation Loss - Training Framework Comparison", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)
    add_figure(doc, FIGURES / "training_val_acc_overlay.png",
               "Figure 14b: Validation Accuracy - Training Framework Comparison", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)

    add_side_by_side(
        doc,
        FIGURES / "lr_1e2_loss.png", FIGURES / "lr_1e2_acc.png",
        "Figure 15a: LR=0.01 Loss", "Figure 15b: LR=0.01 Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "lr_1e3_loss.png", FIGURES / "lr_1e3_acc.png",
        "Figure 16a: LR=0.001 Loss", "Figure 16b: LR=0.001 Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "lr_3e4_loss.png", FIGURES / "lr_3e4_acc.png",
        "Figure 17a: LR=3e-4 Loss", "Figure 17b: LR=3e-4 Accuracy"
    )

    add_heading(doc, "5.5.2 LR Schedulers and Early Stopping", 3)
    rows = [
        ("sched_step", "0.001", "StepLR (step=5, gamma=0.5)", "True", "98.36%", "12"),
        ("sched_plateau", "0.001", "ReduceLROnPlateau (patience=2, factor=0.5)", "True", "98.32%", "18"),
        ("no_early_stop", "0.001", "none", "False", "98.36%", "18"),
    ]
    add_results_table(doc, rows, ["Run", "LR", "Scheduler", "Early Stop", "Best Val Acc", "Best Epoch"])

    add_para(doc, (
        "LR schedulers substantially improve performance. StepLR achieves the best validation accuracy "
        "(98.36%) by epoch 12, also confirmed as the best test accuracy (98.36%). ReduceLROnPlateau reaches "
        "98.32% by epoch 18, adapting the LR dynamically. Training without early stopping also reaches 98.36% "
        "but requires all 20 epochs. These results demonstrate that decaying the learning rate is crucial "
        "for convergence to better minima, and early stopping with StepLR provides an efficient combination."
    ))

    add_side_by_side(
        doc,
        FIGURES / "sched_step_loss.png", FIGURES / "sched_step_acc.png",
        "Figure 18a: StepLR Loss", "Figure 18b: StepLR Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "sched_plateau_loss.png", FIGURES / "sched_plateau_acc.png",
        "Figure 19a: Plateau Loss", "Figure 19b: Plateau Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "no_early_stop_loss.png", FIGURES / "no_early_stop_acc.png",
        "Figure 20a: No Early Stop Loss", "Figure 20b: No Early Stop Accuracy"
    )

    add_figure(doc, FIGURES / "sched_step_confusion.png",
               "Figure 21: Confusion Matrix - Best Model (StepLR)", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)

    doc.add_page_break()

    # --- 5.6 Batch Normalization ---
    add_heading(doc, "5.6 Impact of Batch Normalization", 2)
    add_para(doc, (
        "Batch Normalization (BN) normalizes each layer's pre-activation outputs over the mini-batch, "
        "reducing internal covariate shift. It is applied BEFORE the activation function in this implementation, "
        "following the original paper by Ioffe & Szegedy (2015)."
    ))

    rows = [
        ("bn_off", "No", "97.80%", "5"),
        ("bn_on (baseline)", "Yes", "97.92%", "6"),
    ]
    add_results_table(doc, rows, ["Run", "BatchNorm", "Best Val Acc", "Best Epoch"])

    add_para(doc, (
        "With BatchNorm enabled, the model achieves a higher validation accuracy (97.92% vs 97.80%) "
        "and trains more stably. Without BN, the training and validation loss curves are noisier and "
        "the final accuracy is lower. Placing BN before the activation ensures the activation receives "
        "normalized inputs, maintaining a zero-centered distribution that keeps activations in the linear "
        "regime of ReLU (preventing dead neurons) and ensures gradients flow more uniformly during "
        "backpropagation."
    ))

    add_figure(doc, FIGURES / "batchnorm_val_loss_overlay.png",
               "Figure 22a: Validation Loss - BatchNorm Comparison", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)
    add_figure(doc, FIGURES / "batchnorm_val_acc_overlay.png",
               "Figure 22b: Validation Accuracy - BatchNorm Comparison", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)

    add_side_by_side(
        doc,
        FIGURES / "bn_off_loss.png", FIGURES / "bn_off_acc.png",
        "Figure 23a: BN Off Loss", "Figure 23b: BN Off Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "bn_on_loss.png", FIGURES / "bn_on_acc.png",
        "Figure 24a: BN On Loss", "Figure 24b: BN On Accuracy"
    )

    doc.add_page_break()

    # --- 5.7 Regularization ---
    add_heading(doc, "5.7 Impact of Regularization (L1 / L2)", 2)
    add_para(doc, (
        "L2 regularization (weight decay) and L1 regularization (explicit penalty on absolute weight sum) "
        "were tested at several coefficient values. L1 encourages sparsity, while L2 encourages small weights."
    ))

    rows = [
        ("l2_0 (no reg)", "L2", "0.0", "97.92%", "6"),
        ("l2_1e5", "L2", "1e-5", "97.94%", "6"),
        ("l2_1e4", "L2", "1e-4", "97.82%", "6"),
        ("l2_1e3", "L2", "1e-3", "97.82%", "11"),
        ("l1_1e6", "L1", "1e-6", "97.96%", "6"),
        ("l1_1e5", "L1", "1e-5", "98.02%", "6"),
    ]
    add_results_table(doc, rows, ["Run", "Type", "Lambda", "Best Val Acc", "Best Epoch"])

    add_para(doc, (
        "A very small L2 weight decay (1e-5) provides a marginal improvement over no regularization "
        "(97.94% vs 97.92%). Stronger L2 (1e-4, 1e-3) hurts performance, suggesting over-regularization "
        "for this dataset and model size. Mild L1 regularization (1e-5) achieves the best result in this "
        "group (98.02%), benefiting from the sparsity-inducing property of L1 which effectively selects "
        "the most informative weights. Excessively strong regularization reduces model capacity below what "
        "is needed for MNIST."
    ))

    add_figure(doc, FIGURES / "regularization_val_loss_overlay.png",
               "Figure 25a: Validation Loss - Regularization Comparison", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)
    add_figure(doc, FIGURES / "regularization_val_acc_overlay.png",
               "Figure 25b: Validation Accuracy - Regularization Comparison", width=IMG_OVERLAY_W, height=IMG_OVERLAY_H)

    add_side_by_side(
        doc,
        FIGURES / "l2_0_loss.png", FIGURES / "l2_0_acc.png",
        "Figure 26a: No Reg Loss", "Figure 26b: No Reg Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "l2_1e5_loss.png", FIGURES / "l2_1e5_acc.png",
        "Figure 27a: L2=1e-5 Loss", "Figure 27b: L2=1e-5 Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "l2_1e4_loss.png", FIGURES / "l2_1e4_acc.png",
        "Figure 28a: L2=1e-4 Loss", "Figure 28b: L2=1e-4 Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "l2_1e3_loss.png", FIGURES / "l2_1e3_acc.png",
        "Figure 29a: L2=1e-3 Loss", "Figure 29b: L2=1e-3 Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "l1_1e6_loss.png", FIGURES / "l1_1e6_acc.png",
        "Figure 30a: L1=1e-6 Loss", "Figure 30b: L1=1e-6 Accuracy"
    )
    add_side_by_side(
        doc,
        FIGURES / "l1_1e5_loss.png", FIGURES / "l1_1e5_acc.png",
        "Figure 31a: L1=1e-5 Loss", "Figure 31b: L1=1e-5 Accuracy"
    )

    doc.add_page_break()

    # =========================================================
    # 6. SUMMARY
    # =========================================================
    add_heading(doc, "6. Summary", 1)
    add_para(doc, (
        "The table below summarizes the best validation accuracy achieved across all ablation groups:"
    ))
    rows = [
        ("Baseline", "[256, 128], ReLU, p=0.2, BN, Adam lr=1e-3", "97.92%"),
        ("Best Architecture", "[512, 256, 128] (3 layers)", "98.28%"),
        ("Best Activation", "GELU (marginal gain)", "97.94%"),
        ("Best Dropout", "p=0.5", "98.08%"),
        ("Best Training (LR)", "LR=3e-4, no scheduler", "98.04%"),
        ("Best Training (Scheduler)", "StepLR (step=5, gamma=0.5)", "98.36%"),
        ("Best BatchNorm", "Enabled (before activation)", "97.92%"),
        ("Best L2 Regularization", "lambda=1e-5", "97.94%"),
        ("Best L1 Regularization", "lambda=1e-5", "98.02%"),
        ("OVERALL BEST (Test Acc)", "StepLR scheduler model", "98.36%"),
    ]
    add_results_table(doc, rows, ["Experiment", "Configuration", "Val Acc"])

    add_para(doc, (
        "Key takeaways:"
    ))
    for item in [
        "LR scheduling (StepLR) provides the most impactful improvement, achieving 98.36% test accuracy.",
        "Deeper/wider architectures help on MNIST, with diminishing returns beyond 3 layers.",
        "ReLU and GELU are nearly equivalent for this task.",
        "Dropout and BatchNorm are both beneficial; BN should be placed before the activation.",
        "Mild regularization (L1 lambda=1e-5 or L2 lambda=1e-5) slightly improves generalization.",
        "Early stopping with StepLR offers the best efficiency/performance trade-off.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(11)

    doc.save(str(OUTPUT))
    print(f"Report saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
